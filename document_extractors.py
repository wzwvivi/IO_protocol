# -*- coding: utf-8 -*-
"""
文档抽取模块
负责从 Excel/Word 文件中提取协议内容，转换为统一的中间格式
"""

import os
import re
import json
from typing import Dict, List, Any, Optional, Tuple


class ExtractionResult:
    """文档抽取结果"""
    def __init__(self):
        self.tables: List[Dict[str, Any]] = []  # 抽取的表格
        self.paragraphs: List[str] = []  # 抽取的段落文本
        self.metadata: Dict[str, Any] = {}  # 文档元数据
        self.errors: List[str] = []  # 抽取过程中的错误
        self.warnings: List[str] = []  # 警告信息
        self.raw_text: str = ""  # 原始文本（用于 LLM）
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'tables': self.tables,
            'paragraphs': self.paragraphs,
            'metadata': self.metadata,
            'errors': self.errors,
            'warnings': self.warnings,
            'raw_text': self.raw_text
        }


def extract_excel(file_path: str) -> ExtractionResult:
    """从 Excel 文件抽取内容
    
    Args:
        file_path: Excel 文件路径 (.xlsx, .xls)
    
    Returns:
        ExtractionResult: 抽取结果
    """
    result = ExtractionResult()
    
    try:
        import openpyxl
    except ImportError:
        result.errors.append("缺少 openpyxl 库，请安装: pip install openpyxl")
        return result
    
    if not os.path.exists(file_path):
        result.errors.append(f"文件不存在: {file_path}")
        return result
    
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        result.metadata['filename'] = os.path.basename(file_path)
        result.metadata['sheet_names'] = wb.sheetnames
        result.metadata['file_type'] = 'excel'
        
        all_text_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # 抽取表格数据
            table_data = {
                'sheet_name': sheet_name,
                'rows': [],
                'headers': [],
                'row_count': ws.max_row,
                'col_count': ws.max_column
            }
            
            # 读取所有行
            for row_idx, row in enumerate(ws.iter_rows(max_row=min(500, ws.max_row), values_only=True), 1):
                cells = []
                for cell in row:
                    if cell is not None:
                        cells.append(str(cell).strip())
                    else:
                        cells.append('')
                
                # 跳过全空行
                if not any(c for c in cells):
                    continue
                
                # 第一个非空行可能是表头
                if row_idx == 1 or (not table_data['headers'] and _looks_like_header(cells)):
                    table_data['headers'] = cells
                
                table_data['rows'].append({
                    'row_num': row_idx,
                    'cells': cells
                })
                
                # 拼接文本
                row_text = ' | '.join(c for c in cells if c)
                if row_text:
                    all_text_parts.append(f"[行{row_idx}] {row_text}")
            
            if table_data['rows']:
                result.tables.append(table_data)
        
        # 构建原始文本（供 LLM 使用）
        result.raw_text = '\n'.join(all_text_parts)
        
        wb.close()
        
    except Exception as e:
        result.errors.append(f"读取 Excel 文件失败: {str(e)}")
    
    return result


def extract_word(file_path: str) -> ExtractionResult:
    """从 Word 文件抽取内容
    
    Args:
        file_path: Word 文件路径 (.docx)
    
    Returns:
        ExtractionResult: 抽取结果
    """
    result = ExtractionResult()
    
    try:
        from docx import Document
    except ImportError:
        result.errors.append("缺少 python-docx 库，请安装: pip install python-docx")
        return result
    
    if not os.path.exists(file_path):
        result.errors.append(f"文件不存在: {file_path}")
        return result
    
    try:
        doc = Document(file_path)
        result.metadata['filename'] = os.path.basename(file_path)
        result.metadata['file_type'] = 'word'
        
        all_text_parts = []
        
        # 抽取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                result.paragraphs.append(text)
                all_text_parts.append(text)
        
        # 抽取表格（处理合并单元格）
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'table_index': table_idx,
                'rows': [],
                'headers': [],
                'row_count': len(table.rows),
                'col_count': len(table.columns) if table.rows else 0
            }
            
            seen_row_texts = set()  # 用于行级去重
            
            for row_idx, row in enumerate(table.rows):
                # 处理合并单元格：去除重复的相邻单元格
                cells = []
                prev_cell_text = None
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # 跳过与前一个单元格完全相同的内容（合并单元格的重复）
                    if cell_text != prev_cell_text:
                        cells.append(cell_text)
                        prev_cell_text = cell_text
                
                # 跳过全空行
                if not any(c for c in cells):
                    continue
                
                # 行级去重：跳过完全相同的行
                row_signature = '|'.join(cells)
                if row_signature in seen_row_texts:
                    continue
                seen_row_texts.add(row_signature)
                
                # 第一行可能是表头
                if row_idx == 0 or (not table_data['headers'] and _looks_like_header(cells)):
                    table_data['headers'] = cells
                
                table_data['rows'].append({
                    'row_num': row_idx + 1,
                    'cells': cells
                })
                
                # 拼接文本
                row_text = ' | '.join(c for c in cells if c)
                if row_text:
                    all_text_parts.append(f"[表{table_idx + 1}行{row_idx + 1}] {row_text}")
            
            if table_data['rows']:
                result.tables.append(table_data)
        
        # 构建原始文本
        result.raw_text = '\n'.join(all_text_parts)
        
    except Exception as e:
        result.errors.append(f"读取 Word 文件失败: {str(e)}")
    
    return result


def _looks_like_header(cells: List[str]) -> bool:
    """判断一行是否看起来像表头"""
    if not cells:
        return False
    
    # 表头通常包含这些关键词
    header_keywords = [
        'label', 'Label', 'LABEL', '标签', 'bit', 'Bit', 'BIT',
        '名称', 'name', 'Name', '方向', 'direction', '说明', 'description',
        '分辨率', 'resolution', '单位', 'unit', '范围', 'range',
        '数据类型', 'type', '备注', 'notes', '来源', 'source',
        '发送', '接收', 'TX', 'RX', 'SDI', 'SSM', '定义'
    ]
    
    # 检查是否包含关键词
    for cell in cells:
        for keyword in header_keywords:
            if keyword in cell:
                return True
    
    return False


def extract_document(file_path: str) -> ExtractionResult:
    """根据文件类型自动选择抽取方法
    
    Args:
        file_path: 文件路径
    
    Returns:
        ExtractionResult: 抽取结果
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in ['.xlsx', '.xls']:
        return extract_excel(file_path)
    elif ext == '.docx':
        return extract_word(file_path)
    else:
        result = ExtractionResult()
        result.errors.append(f"不支持的文件类型: {ext}")
        return result


# ============================================================
# 规则预解析：识别 ARINC429 协议相关字段
# ============================================================

def preparse_arinc429_table(table: Dict[str, Any]) -> Dict[str, Any]:
    """对表格进行 ARINC429 协议预解析
    
    尝试识别：
    - Label 列（八进制/十进制）
    - 名称/信号名
    - 方向（发送/接收）
    - 位定义（BNR/离散/枚举）
    - 分辨率、单位、范围
    
    Args:
        table: 表格数据
    
    Returns:
        解析结果，包含识别出的字段映射和初步解析的 Labels
    """
    result = {
        'column_mapping': {},  # 列名 -> 标准字段名
        'parsed_labels': [],   # 初步解析的 Labels
        'confidence': 0.0,     # 置信度
        'issues': []           # 问题
    }
    
    headers = table.get('headers', [])
    rows = table.get('rows', [])
    
    if not headers or not rows:
        return result
    
    # 识别列映射
    column_mapping = _identify_columns(headers)
    result['column_mapping'] = column_mapping
    
    # 如果没有识别出 Label 列，置信度很低
    if 'label' not in column_mapping:
        result['issues'].append("未识别出 Label 列")
        result['confidence'] = 0.1
        return result
    
    # 解析每一行
    label_col_idx = column_mapping['label']
    name_col_idx = column_mapping.get('name')
    direction_col_idx = column_mapping.get('direction')
    
    for row in rows:
        cells = row.get('cells', [])
        
        # 跳过表头行
        if cells == headers:
            continue
        
        # 提取 Label
        label_val = cells[label_col_idx] if label_col_idx < len(cells) else ''
        if not label_val:
            continue
        
        # 尝试解析为八进制
        label_oct = _parse_label_oct(label_val)
        if not label_oct:
            result['issues'].append(f"无法解析 Label: {label_val}")
            continue
        
        # 提取更多字段
        resolution_col_idx = column_mapping.get('resolution')
        unit_col_idx = column_mapping.get('unit')
        range_col_idx = column_mapping.get('range')
        
        # 尝试提取分辨率
        resolution = None
        if resolution_col_idx is not None and resolution_col_idx < len(cells):
            res_str = cells[resolution_col_idx]
            resolution = _parse_resolution(res_str)
        
        # 尝试提取单位
        unit = ''
        if unit_col_idx is not None and unit_col_idx < len(cells):
            unit = cells[unit_col_idx].strip()
        
        # 尝试提取范围
        range_str = ''
        if range_col_idx is not None and range_col_idx < len(cells):
            range_str = cells[range_col_idx].strip()
        
        # 构建 Label 结构
        label = {
            'label_oct': label_oct,
            'name': cells[name_col_idx] if name_col_idx is not None and name_col_idx < len(cells) else '',
            'direction': cells[direction_col_idx] if direction_col_idx is not None and direction_col_idx < len(cells) else '',
            'sources': [],
            'discrete_bits': {},
            'special_fields': [],
            'bnr_fields': [],
            'notes': '',
            '_raw_row': cells,  # 保留原始行数据供 LLM 进一步解析
            '_confidence': 0.5,  # 初步置信度
            # 保存提取的额外信息供 LLM 参考
            '_extracted_resolution': resolution,
            '_extracted_unit': unit,
            '_extracted_range': range_str
        }
        
        result['parsed_labels'].append(label)
    
    # 计算整体置信度
    if result['parsed_labels']:
        result['confidence'] = 0.5 + (0.1 * len(column_mapping))
        result['confidence'] = min(result['confidence'], 0.9)
    
    return result


def _identify_columns(headers: List[str]) -> Dict[str, int]:
    """识别列名对应的标准字段
    
    Args:
        headers: 表头列表
    
    Returns:
        字段名 -> 列索引 的映射
    """
    mapping = {}
    
    label_patterns = [
        r'label', r'Label', r'LABEL', r'标签', r'地址', r'编号'
    ]
    name_patterns = [
        r'名称', r'name', r'Name', r'NAME', r'信号名', r'参数名'
    ]
    direction_patterns = [
        r'方向', r'direction', r'Direction', r'收发', r'TX/RX', r'发送|接收'
    ]
    bit_patterns = [
        r'bit', r'Bit', r'BIT', r'位', r'位号', r'位定义'
    ]
    resolution_patterns = [
        r'分辨率', r'resolution', r'Resolution', r'精度', r'LSB'
    ]
    unit_patterns = [
        r'单位', r'unit', r'Unit', r'UNIT'
    ]
    range_patterns = [
        r'范围', r'range', r'Range', r'RANGE', r'量程'
    ]
    notes_patterns = [
        r'备注', r'notes', r'Notes', r'说明', r'描述', r'description'
    ]
    
    for idx, header in enumerate(headers):
        header_lower = header.lower()
        
        # Label
        if not mapping.get('label'):
            for pattern in label_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['label'] = idx
                    break
        
        # Name
        if not mapping.get('name'):
            for pattern in name_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['name'] = idx
                    break
        
        # Direction
        if not mapping.get('direction'):
            for pattern in direction_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['direction'] = idx
                    break
        
        # Bit
        if not mapping.get('bit'):
            for pattern in bit_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['bit'] = idx
                    break
        
        # Resolution
        if not mapping.get('resolution'):
            for pattern in resolution_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['resolution'] = idx
                    break
        
        # Unit
        if not mapping.get('unit'):
            for pattern in unit_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['unit'] = idx
                    break
        
        # Range
        if not mapping.get('range'):
            for pattern in range_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['range'] = idx
                    break
        
        # Notes
        if not mapping.get('notes'):
            for pattern in notes_patterns:
                if re.search(pattern, header, re.IGNORECASE):
                    mapping['notes'] = idx
                    break
    
    return mapping


def _parse_resolution(value: str) -> Optional[float]:
    """解析分辨率值
    
    支持：
    - 纯数字: 0.0055, 0.125
    - 分数: 1/2^14, 1/16384
    - 科学计数法: 1e-4, 5.5e-3
    
    Args:
        value: 分辨率原始值
    
    Returns:
        浮点数分辨率，解析失败返回 None
    """
    if not value:
        return None
    
    value = str(value).strip()
    
    # 移除单位
    value = re.sub(r'[°度%].*$', '', value).strip()
    
    try:
        # 直接解析为浮点数
        return float(value)
    except ValueError:
        pass
    
    # 尝试解析分数形式 1/2^14 或 1/16384
    frac_match = re.match(r'(\d+(?:\.\d+)?)\s*/\s*(\d+(?:\.\d+)?|\d+\^(\d+))', value)
    if frac_match:
        numerator = float(frac_match.group(1))
        if frac_match.group(3):
            # 2^14 形式
            base = 2
            exp = int(frac_match.group(3))
            denominator = base ** exp
        else:
            denominator = float(frac_match.group(2))
        
        if denominator != 0:
            return numerator / denominator
    
    # 尝试解析 2^-14 形式
    pow_match = re.match(r'2\^(-?\d+)', value)
    if pow_match:
        exp = int(pow_match.group(1))
        return 2 ** exp
    
    return None


def _parse_label_oct(value: str) -> Optional[str]:
    """解析 Label 值为八进制字符串
    
    支持：
    - 八进制: 115, 0o115, 115(o), 115o
    - 十进制（带标注）: 77(d), 77d
    - 十六进制（带标注）: 0x4D, 4Dh
    
    Args:
        value: Label 原始值
    
    Returns:
        三位八进制字符串，如 "115"，解析失败返回 None
    """
    value = str(value).strip()
    
    if not value:
        return None
    
    # 移除常见前缀/后缀
    value = value.replace('Label', '').replace('label', '').replace('LABEL', '')
    value = value.strip()
    
    # 尝试解析为八进制
    try:
        # 0o 前缀
        if value.lower().startswith('0o'):
            oct_val = int(value, 8)
            return f"{oct_val:03o}"
        
        # (o) 或 o 后缀
        if value.lower().endswith('o') or value.lower().endswith('(o)'):
            num = re.sub(r'[oO()\s]', '', value)
            oct_val = int(num, 8)
            return f"{oct_val:03o}"
        
        # (d) 或 d 后缀 - 十进制
        if value.lower().endswith('d') or value.lower().endswith('(d)'):
            num = re.sub(r'[dD()\s]', '', value)
            dec_val = int(num, 10)
            return f"{dec_val:03o}"
        
        # 0x 前缀 - 十六进制
        if value.lower().startswith('0x'):
            hex_val = int(value, 16)
            return f"{hex_val:03o}"
        
        # h 后缀 - 十六进制
        if value.lower().endswith('h'):
            num = value[:-1]
            hex_val = int(num, 16)
            return f"{hex_val:03o}"
        
        # 纯数字 - 假定为八进制
        if re.match(r'^[0-7]+$', value):
            oct_val = int(value, 8)
            return f"{oct_val:03o}"
        
        # 纯数字（包含 8/9）- 可能是十进制
        if re.match(r'^\d+$', value):
            dec_val = int(value, 10)
            # 如果数值在合理范围内（0-255），转为八进制
            if 0 <= dec_val <= 255:
                return f"{dec_val:03o}"
        
    except (ValueError, TypeError):
        pass
    
    return None


def build_llm_prompt_context(extraction_result: ExtractionResult, preparse_results: List[Dict], max_chars: int = 25000) -> str:
    """构建供 LLM 使用的上下文
    
    Args:
        extraction_result: 文档抽取结果
        preparse_results: 预解析结果列表
        max_chars: 最大字符数限制（默认 25000，约 6000-8000 tokens）
    
    Returns:
        格式化的上下文字符串
    """
    parts = []
    current_chars = 0
    
    # 文档元信息
    header = f"# 文档信息\n文件名: {extraction_result.metadata.get('filename', '未知')}\n文件类型: {extraction_result.metadata.get('file_type', '未知')}\n"
    parts.append(header)
    current_chars += len(header)
    
    # 智能选择表格：优先选择包含 Label 定义的表格
    tables_with_labels = []
    tables_without_labels = []
    
    for i, table in enumerate(extraction_result.tables):
        # 检查表格是否可能包含 Label 定义
        has_label_info = False
        for row in table.get('rows', [])[:10]:
            cells = row.get('cells', [])
            row_text = ' '.join(str(c) for c in cells).lower()
            if any(kw in row_text for kw in ['label', '标号', '规范号', 'bit', '位', '分辨率', 'resolution']):
                has_label_info = True
                break
        
        if has_label_info:
            tables_with_labels.append((i, table))
        else:
            tables_without_labels.append((i, table))
    
    # 优先处理包含 Label 的表格
    selected_tables = tables_with_labels[:15] + tables_without_labels[:5]  # 最多 20 个表格
    
    for idx, table in selected_tables:
        if current_chars >= max_chars:
            parts.append(f"\n... (内容已截断，共 {len(extraction_result.tables)} 个表格)")
            break
        
        table_parts = []
        table_parts.append(f"\n## 表格 {idx + 1}")
        if table.get('sheet_name'):
            table_parts.append(f"工作表: {table['sheet_name']}")
        
        # 表头
        if table.get('headers'):
            headers_str = ' | '.join(str(h)[:30] for h in table['headers'][:10])  # 限制表头
            table_parts.append(f"表头: {headers_str}")
        
        # 数据行（限制数量和每行长度）
        table_parts.append("数据:")
        rows = table.get('rows', [])
        max_rows = min(25, len(rows))  # 每个表格最多 25 行
        
        for row in rows[:max_rows]:
            cells = row.get('cells', [])
            # 限制每个单元格长度和列数
            cells_str = ' | '.join(str(c)[:50] for c in cells[:10])
            if len(cells_str) > 200:
                cells_str = cells_str[:200] + '...'
            table_parts.append(f"  {cells_str}")
        
        if len(rows) > max_rows:
            table_parts.append(f"  ... (共 {len(rows)} 行)")
        
        table_text = '\n'.join(table_parts)
        
        # 检查是否超出限制
        if current_chars + len(table_text) > max_chars:
            # 尝试添加截断版本
            remaining = max_chars - current_chars - 100
            if remaining > 500:
                parts.append(table_text[:remaining] + "\n... (表格内容已截断)")
            break
        
        parts.append(table_text)
        current_chars += len(table_text)
    
    # 如果还有空间，添加一些段落
    if current_chars < max_chars - 1000 and extraction_result.paragraphs:
        parts.append("\n## 关键段落")
        for para in extraction_result.paragraphs[:10]:
            if current_chars >= max_chars:
                break
            para_text = para[:200] if len(para) > 200 else para
            parts.append(f"  {para_text}")
            current_chars += len(para_text)
    
    result = '\n'.join(parts)
    
    # 最终检查
    if len(result) > max_chars:
        result = result[:max_chars] + "\n... (内容已截断)"
    
    return result


def convert_to_markdown(extraction_result: ExtractionResult, max_chars: int = 25000) -> str:
    """将抽取结果转换为紧凑的 Markdown 格式
    
    Args:
        extraction_result: 文档抽取结果
        max_chars: 最大字符数限制
    
    Returns:
        Markdown 格式的文档内容
    """
    md_parts = []
    current_chars = 0
    
    # 文档标题
    filename = extraction_result.metadata.get('filename', '未知文档')
    md_parts.append(f"# {filename}\n")
    current_chars += len(md_parts[-1])
    
    # 智能选择和排序表格
    # 优先处理包含 Label 定义的表格
    label_tables = []
    other_tables = []
    
    for i, table in enumerate(extraction_result.tables):
        # 检查是否是 Label 定义表格
        is_label_table = False
        table_text = ''
        for row in table.get('rows', [])[:5]:
            cells = row.get('cells', [])
            row_text = ' '.join(str(c) for c in cells).lower()
            table_text += row_text
        
        if any(kw in table_text for kw in ['label', '标号', '规范号', 'bit', '位定义', '数据格式']):
            is_label_table = True
        
        if is_label_table:
            label_tables.append((i, table))
        else:
            other_tables.append((i, table))
    
    # 合并：优先 Label 表格，增加数量限制
    # ATG 文档每个 Label 一个表格，可能有 30+ 个 Label
    all_tables = label_tables[:50] + other_tables[:3]  # 最多 50 个 Label 表格 + 3 个其他表格
    
    # 统一单元格规范化，便于去重
    def _normalize_cell_text(value: Any) -> str:
        text = str(value or '')
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    # 转换表格为 Markdown
    for idx, table in all_tables:
        if current_chars >= max_chars - 500:
            md_parts.append(f"\n> ⚠️ 内容已截断，共 {len(extraction_result.tables)} 个表格\n")
            break
        
        # 表格标题
        sheet_name = table.get('sheet_name', '')
        if sheet_name:
            md_parts.append(f"\n## 表格 {idx + 1}: {sheet_name}\n")
        else:
            md_parts.append(f"\n## 表格 {idx + 1}\n")
        
        rows = table.get('rows', [])
        if not rows:
            continue
        
        # 确定列数（取前几行的最大列数）
        max_cols = 0
        for row in rows[:10]:
            cells = row.get('cells', [])
            max_cols = max(max_cols, len(cells))
        max_cols = min(max_cols, 12)  # 最多 12 列
        
        # 生成 Markdown 表格
        # 表头（不截断）
        headers = table.get('headers', [])
        if headers:
            header_cells = [str(h) for h in headers[:max_cols]]
        else:
            # 使用第一行作为表头
            if rows:
                header_cells = [str(c) for c in rows[0].get('cells', [])[:max_cols]]
            else:
                header_cells = [f'列{i+1}' for i in range(max_cols)]
        
        # 补齐列数
        while len(header_cells) < max_cols:
            header_cells.append('')
        
        md_parts.append('| ' + ' | '.join(header_cells) + ' |')
        md_parts.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
        
        # 数据行（增加行数限制，确保位定义完整）
        start_row = 1 if not headers else 0
        max_rows = min(50, len(rows))  # 每个表格最多 50 行，确保完整的位定义
        
        seen_row_signatures = set()
        for row in rows[start_row:max_rows]:
            cells = row.get('cells', [])

            # 行去重：Word 合并单元格展开后常出现重复行
            signature_cells = [_normalize_cell_text(c) for c in cells[:max_cols]]
            row_signature = '||'.join(signature_cells)
            if row_signature in seen_row_signatures:
                continue
            seen_row_signatures.add(row_signature)

            # 清理单元格内容（不截断）
            row_cells = []
            for c in cells[:max_cols]:
                cell_text = _normalize_cell_text(c).replace('|', '/')  # 只替换管道符，不截断
                row_cells.append(cell_text)
            
            # 补齐列数
            while len(row_cells) < max_cols:
                row_cells.append('')
            
            md_parts.append('| ' + ' | '.join(row_cells) + ' |')
        
        if len(rows) > max_rows:
            md_parts.append(f'\n> 表格共 {len(rows)} 行，已截断显示前 {max_rows} 行\n')
        
        md_parts.append('')  # 空行分隔
        
        # 更新字符计数
        current_chars = sum(len(p) for p in md_parts)
    
    # 添加关键段落（如果有空间）
    if current_chars < max_chars - 2000 and extraction_result.paragraphs:
        md_parts.append("\n## 文档说明\n")
        
        # 筛选有意义的段落
        meaningful_paras = []
        for para in extraction_result.paragraphs:
            para = para.strip()
            if len(para) > 10 and not para.startswith('表'):
                meaningful_paras.append(para)
        
        for para in meaningful_paras[:15]:
            if current_chars >= max_chars:
                break
            md_parts.append(f"- {para}")  # 不截断段落
            current_chars += len(para)
    
    result = '\n'.join(md_parts)
    
    # 最终长度检查
    if len(result) > max_chars:
        result = result[:max_chars] + "\n\n> ⚠️ 内容已截断"
    
    return result


def parse_bit_definition_table(table: Dict[str, Any]) -> Dict[str, Any]:
    """解析位定义表，提取结构化的位定义数据
    
    Args:
        table: 表格数据
    
    Returns:
        结构化的 Label 定义，包含位定义
    """
    rows = table.get('rows', [])
    if not rows:
        return None
    
    result = {
        'label_oct': None,
        'name': None,
        'direction': None,
        'source': None,
        'dest': None,
        'bit_definitions': {}  # {位号: 说明}
    }
    
    # 解析表格行
    for row in rows:
        cells = row.get('cells', [])
        if len(cells) < 2:
            continue
        
        first_cell = str(cells[0]).strip() if cells[0] else ''
        
        # 提取 Label 号
        if '标' in first_cell and '号' in first_cell:
            for cell in cells[1:]:
                cell_str = str(cell).strip()
                if cell_str.isdigit() and len(cell_str) <= 3:
                    result['label_oct'] = cell_str
                    break
        
        # 提取信号名称
        if '信号名称' in first_cell or '信号名' in first_cell:
            for cell in cells[1:]:
                cell_str = str(cell).strip()
                if cell_str and cell_str not in ['', 'N/A']:
                    result['name'] = cell_str
                    break
        
        # 提取源
        if first_cell == '数据内容' and len(cells) > 2:
            second_cell = str(cells[1]).strip() if cells[1] else ''
            if '源' in second_cell:
                for cell in cells[2:]:
                    cell_str = str(cell).strip()
                    if cell_str and cell_str not in ['', 'N/A', '单   位', '单位']:
                        result['source'] = cell_str
                        break
            elif '目' in second_cell:
                for cell in cells[2:]:
                    cell_str = str(cell).strip()
                    if cell_str and cell_str not in ['', 'N/A', '信号范围']:
                        result['dest'] = cell_str
                        break
        
        # 提取位定义（数据格式行）
        if first_cell == '数据格式' and len(cells) >= 3:
            # 每行可能有多组 (位号, 说明)
            # 格式: [数据格式, 位号1, 说明1, 位号2, 说明2, 位号3, 说明3]
            i = 1
            while i < len(cells) - 1:
                bit_num_str = str(cells[i]).strip() if cells[i] else ''
                desc_str = str(cells[i + 1]).strip() if i + 1 < len(cells) and cells[i + 1] else ''
                
                # 检查是否是有效的位号
                if bit_num_str.isdigit():
                    bit_num = int(bit_num_str)
                    if 1 <= bit_num <= 32 and desc_str:
                        # 跳过 Label 编码位 (1-8) 的 0/1 值
                        if bit_num <= 8 and desc_str in ['0', '1']:
                            pass
                        # 跳过 SDI, SSM, 校验位
                        elif desc_str.upper() in ['SDI', 'SSM', '校验位', '校验位：奇校验', '奇校验']:
                            pass
                        else:
                            result['bit_definitions'][bit_num] = desc_str
                
                i += 2  # 移动到下一组
    
    # 构建方向
    if result['source'] and result['dest']:
        result['direction'] = f"{result['source']} -> {result['dest']}"
    
    return result if result['label_oct'] else None


def format_structured_label_data(parsed_labels: List[Dict]) -> str:
    """将解析的 Label 数据格式化为结构化文本，供 LLM 使用
    
    Args:
        parsed_labels: 解析的 Label 列表
    
    Returns:
        结构化的文本
    """
    if not parsed_labels:
        return ""
    
    parts = ["\n## 结构化位定义数据\n"]
    parts.append("以下是从文档中提取的结构化位定义，请基于此数据生成 JSON 输出：\n")
    
    for label_data in parsed_labels:
        if not label_data:
            continue
        
        label_oct = label_data.get('label_oct', '?')
        name = label_data.get('name', '')
        direction = label_data.get('direction', '')
        source = label_data.get('source', '')
        bit_defs = label_data.get('bit_definitions', {})
        
        parts.append(f"### Label {label_oct}: {name}")
        if direction:
            parts.append(f"- 方向: {direction}")
        if source:
            parts.append(f"- 源: {source}")
        
        if bit_defs:
            parts.append("- 位定义:")
            # 按位号排序
            for bit_num in sorted(bit_defs.keys()):
                desc = bit_defs[bit_num]
                parts.append(f"  - Bit {bit_num}: {desc}")
        
        parts.append("")
    
    return '\n'.join(parts)


def build_llm_prompt_context_v2(extraction_result: ExtractionResult, preparse_results: List[Dict] = None, max_chars: int = 100000) -> str:
    """构建供 LLM 使用的上下文（V2 版本，使用结构化数据）
    
    Args:
        extraction_result: 文档抽取结果
        preparse_results: 预解析结果列表（可选）
        max_chars: 最大字符数限制
    
    Returns:
        结构化的上下文字符串
    """
    parts = []
    
    # 1. 解析所有表格，提取结构化位定义
    structured_labels = []
    for table in extraction_result.tables:
        parsed = parse_bit_definition_table(table)
        if parsed and parsed.get('label_oct'):
            # 去重
            existing = [l for l in structured_labels if l.get('label_oct') == parsed.get('label_oct')]
            if not existing:
                structured_labels.append(parsed)
    
    # 2. 如果有结构化数据，优先使用
    if structured_labels:
        parts.append(format_structured_label_data(structured_labels))
        parts.append(f"\n共解析出 {len(structured_labels)} 个 Label（结构化数据）\n")
    
    # 3. 始终补充 Markdown 表格，确保不遗漏任何 Label
    # 因为 parse_bit_definition_table 可能无法解析所有表格格式
    remaining_chars = max_chars - len('\n'.join(parts)) - 1000
    if remaining_chars > 5000:
        md_content = convert_to_markdown(extraction_result, remaining_chars)
        parts.append("\n## 原始表格数据（完整）\n")
        parts.append(md_content)
    
    result = '\n'.join(parts)
    
    # 最终长度检查
    if len(result) > max_chars:
        result = result[:max_chars] + "\n\n> ⚠️ 内容已截断"
    
    return result


if __name__ == '__main__':
    # 测试
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"抽取文件: {file_path}")
        result = extract_document(file_path)
        print(f"表格数: {len(result.tables)}")
        print(f"段落数: {len(result.paragraphs)}")
        print(f"错误: {result.errors}")
        print(f"警告: {result.warnings}")
        
        # 预解析
        for i, table in enumerate(result.tables):
            print(f"\n表格 {i + 1} 预解析:")
            preparse = preparse_arinc429_table(table)
            print(f"  列映射: {preparse['column_mapping']}")
            print(f"  解析出 {len(preparse['parsed_labels'])} 个 Label")
            print(f"  置信度: {preparse['confidence']:.1%}")
        
        # 测试 Markdown 转换
        print("\n" + "=" * 50)
        print("Markdown 转换结果:")
        print("=" * 50)
        md = convert_to_markdown(result, 5000)
        print(md[:3000])
        print(f"\n... (总长度: {len(md)} 字符)")